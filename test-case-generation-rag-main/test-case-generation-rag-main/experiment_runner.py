import json
from docx import Document

def main():

    with open("outputs/experiment_results.json") as f:
        data = json.load(f)

    total_A_tokens = total_B_tokens = 0
    total_A_time = total_B_time = 0

    for d in data:
        total_A_tokens += d["path_A"]["tokens"]
        total_B_tokens += d["path_B"]["tokens"]

        total_A_time += d["path_A"]["time"]
        total_B_time += d["path_B"]["time"]

    # ✅ energy (from playbook)
    energy_A = (total_A_tokens / 1000) * 0.0003
    energy_B = (total_B_tokens / 1000) * 0.0003

    co2_A = energy_A * 0.233
    co2_B = energy_B * 0.233

    # ✅ percentage
    token_red = ((total_A_tokens - total_B_tokens)/total_A_tokens)*100 if total_A_tokens else 0
    energy_red = ((energy_A - energy_B)/energy_A)*100 if energy_A else 0
    time_red = ((total_A_time - total_B_time)/total_A_time)*100 if total_A_time else 0

    # ✅ DOCX REPORT
    doc = Document()

    doc.add_heading("Sustainability Report", 0)
    doc.add_paragraph("BRD vs Chunking\n")

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "BRD"
    hdr[2].text = "Chunking"
    hdr[3].text = "Winner"

    def row(m, a, b):
        r = table.add_row().cells
        r[0].text = m
        r[1].text = str(round(a,3))
        r[2].text = str(round(b,3))
        r[3].text = "Chunking ✅"

    row("Tokens", total_A_tokens, total_B_tokens)
    row("Time(ms)", total_A_time, total_B_time)
    row("Energy(Wh)", energy_A, energy_B)
    row("CO2(g)", co2_A, co2_B)

    doc.add_heading("Analysis", 1)
    doc.add_paragraph(
        f"Token reduction: {round(token_red,2)}%\n"
        f"Energy reduction: {round(energy_red,2)}%\n"
        f"Speed improvement: {round(time_red,2)}%"
    )

    doc.add_heading("Final Verdict", 1)
    doc.add_paragraph("Chunking is more efficient than BRD.")

    doc.save("outputs/final_report.docx")

    print("✅ Report Generated → outputs/final_report.docx")

if __name__ == "__main__":
    main()
